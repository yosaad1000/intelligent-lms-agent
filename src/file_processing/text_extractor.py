"""
Advanced text extraction utilities with AWS Textract integration
Supports PDF, DOCX, TXT files, and images using AWS Textract for enhanced extraction
"""

import io
import logging
import boto3
import json
import time
from typing import Optional, Dict, Any, List
import mimetypes
import base64

# Text extraction libraries
try:
    import PyPDF2
    from docx import Document
    TEXT_EXTRACTION_AVAILABLE = True
except ImportError:
    TEXT_EXTRACTION_AVAILABLE = False
    print("Warning: Text extraction libraries not available. Install with: pip install PyPDF2 python-docx")

logger = logging.getLogger(__name__)


class TextExtractor:
    """Advanced text extraction utility with AWS Textract integration"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
    SUPPORTED_MIME_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'image/png',
        'image/jpeg',
        'image/tiff',
        'image/bmp'
    }
    
    # Textract supported formats
    TEXTRACT_SUPPORTED = {'.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
    
    def __init__(self):
        self.extraction_methods = {
            '.pdf': self._extract_from_pdf_textract,
            '.docx': self._extract_from_docx,
            '.txt': self._extract_from_txt,
            '.png': self._extract_from_image_textract,
            '.jpg': self._extract_from_image_textract,
            '.jpeg': self._extract_from_image_textract,
            '.tiff': self._extract_from_image_textract,
            '.bmp': self._extract_from_image_textract
        }
        
        # Initialize AWS clients
        self.textract_client = boto3.client('textract')
        self.comprehend_client = boto3.client('comprehend')
        
        # Textract configuration
        self.use_textract = True  # Can be disabled for fallback
        self.textract_timeout = 300  # 5 minutes for async jobs
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported for text extraction"""
        
        file_ext = self._get_file_extension(filename)
        return file_ext in self.SUPPORTED_EXTENSIONS
    
    def extract_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from file content
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename to determine file type
            
        Returns:
            Dictionary with extraction results
        """
        
        try:
            file_ext = self._get_file_extension(filename)
            
            if not self.is_supported_file(filename):
                return {
                    'success': False,
                    'error': f'Unsupported file type: {file_ext}',
                    'text': '',
                    'metadata': {}
                }
            
            if not TEXT_EXTRACTION_AVAILABLE and file_ext in ['.pdf', '.docx']:
                return {
                    'success': False,
                    'error': 'Text extraction libraries not available',
                    'text': '',
                    'metadata': {}
                }
            
            # Extract text using appropriate method
            extraction_method = self.extraction_methods[file_ext]
            result = extraction_method(file_content, filename)
            
            # Add metadata
            result['metadata'] = {
                'file_extension': file_ext,
                'file_size': len(file_content),
                'text_length': len(result.get('text', '')),
                'extraction_method': extraction_method.__name__
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension in lowercase with dot"""
        import os
        return os.path.splitext(filename.lower())[1]
    
    def _extract_from_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return {
                        'success': True,
                        'text': text,
                        'encoding_used': encoding
                    }
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            return {
                'success': True,
                'text': text,
                'encoding_used': 'utf-8 (with errors replaced)',
                'warning': 'Some characters may have been replaced due to encoding issues'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from TXT file: {str(e)}',
                'text': ''
            }
    
    def _extract_from_pdf_textract(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using AWS Textract with fallback to PyPDF2"""
        
        if self.use_textract:
            try:
                # Try Textract first for better accuracy
                textract_result = self._extract_with_textract(file_content, filename, 'PDF')
                if textract_result['success']:
                    return textract_result
                else:
                    logger.warning(f"Textract failed for {filename}, falling back to PyPDF2: {textract_result['error']}")
            except Exception as e:
                logger.warning(f"Textract error for {filename}, falling back to PyPDF2: {str(e)}")
        
        # Fallback to PyPDF2
        return self._extract_from_pdf_pypdf2(file_content, filename)
    
    def _extract_from_pdf_pypdf2(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF file using PyPDF2 (fallback method)"""
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return {
                    'success': False,
                    'error': 'PDF is encrypted and cannot be processed',
                    'text': ''
                }
            
            # Extract text from all pages
            text_content = []
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            full_text = '\n\n'.join(text_content)
            
            return {
                'success': True,
                'text': full_text,
                'page_count': page_count,
                'pages_processed': len(text_content),
                'extraction_method': 'PyPDF2'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from PDF: {str(e)}',
                'text': ''
            }
    
    def _extract_from_image_textract(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from image using AWS Textract"""
        
        try:
            return self._extract_with_textract(file_content, filename, 'IMAGE')
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from image: {str(e)}',
                'text': ''
            }
    
    def _extract_with_textract(self, file_content: bytes, filename: str, document_type: str) -> Dict[str, Any]:
        """Extract text using AWS Textract (supports PDF and images)"""
        
        try:
            # Check file size (Textract limits: 10MB for sync, 500MB for async)
            file_size = len(file_content)
            max_sync_size = 10 * 1024 * 1024  # 10MB
            
            if file_size > max_sync_size:
                # Use async processing for large files
                return self._extract_with_textract_async(file_content, filename, document_type)
            else:
                # Use synchronous processing for smaller files
                return self._extract_with_textract_sync(file_content, filename, document_type)
                
        except Exception as e:
            logger.error(f"Textract extraction error for {filename}: {str(e)}")
            return {
                'success': False,
                'error': f'Textract processing failed: {str(e)}',
                'text': ''
            }
    
    def _extract_with_textract_sync(self, file_content: bytes, filename: str, document_type: str) -> Dict[str, Any]:
        """Synchronous Textract extraction for smaller files"""
        
        try:
            # Call Textract detect_document_text
            response = self.textract_client.detect_document_text(
                Document={'Bytes': file_content}
            )
            
            # Extract text from response
            text_blocks = []
            line_blocks = []
            word_blocks = []
            
            for block in response.get('Blocks', []):
                if block['BlockType'] == 'LINE':
                    line_blocks.append(block['Text'])
                elif block['BlockType'] == 'WORD':
                    word_blocks.append(block['Text'])
            
            # Combine lines to form full text
            full_text = '\n'.join(line_blocks)
            
            # Get additional analysis with Comprehend
            comprehend_analysis = self._analyze_with_comprehend(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'extraction_method': 'AWS Textract (sync)',
                'document_type': document_type,
                'blocks_detected': len(response.get('Blocks', [])),
                'lines_detected': len(line_blocks),
                'words_detected': len(word_blocks),
                'comprehend_analysis': comprehend_analysis
            }
            
        except Exception as e:
            logger.error(f"Textract sync extraction failed: {str(e)}")
            return {
                'success': False,
                'error': f'Textract sync processing failed: {str(e)}',
                'text': ''
            }
    
    def _extract_with_textract_async(self, file_content: bytes, filename: str, document_type: str) -> Dict[str, Any]:
        """Asynchronous Textract extraction for larger files"""
        
        try:
            # For async processing, we need to upload to S3 first
            # This is a simplified version - in production, you'd want proper S3 handling
            
            # Start async job
            response = self.textract_client.start_document_text_detection(
                DocumentLocation={
                    'S3Object': {
                        'Bucket': 'temp-textract-bucket',  # Would need to be configured
                        'Name': f'temp/{filename}'
                    }
                }
            )
            
            job_id = response['JobId']
            
            # Poll for completion
            max_wait_time = self.textract_timeout
            wait_interval = 5
            elapsed_time = 0
            
            while elapsed_time < max_wait_time:
                result = self.textract_client.get_document_text_detection(JobId=job_id)
                status = result['JobStatus']
                
                if status == 'SUCCEEDED':
                    # Extract text from results
                    text_blocks = []
                    
                    for block in result.get('Blocks', []):
                        if block['BlockType'] == 'LINE':
                            text_blocks.append(block['Text'])
                    
                    full_text = '\n'.join(text_blocks)
                    
                    # Get Comprehend analysis
                    comprehend_analysis = self._analyze_with_comprehend(full_text)
                    
                    return {
                        'success': True,
                        'text': full_text,
                        'extraction_method': 'AWS Textract (async)',
                        'document_type': document_type,
                        'job_id': job_id,
                        'processing_time': elapsed_time,
                        'comprehend_analysis': comprehend_analysis
                    }
                    
                elif status == 'FAILED':
                    return {
                        'success': False,
                        'error': f'Textract async job failed: {result.get("StatusMessage", "Unknown error")}',
                        'text': ''
                    }
                
                # Wait before next poll
                time.sleep(wait_interval)
                elapsed_time += wait_interval
            
            # Timeout reached
            return {
                'success': False,
                'error': f'Textract async job timed out after {max_wait_time} seconds',
                'text': ''
            }
            
        except Exception as e:
            logger.error(f"Textract async extraction failed: {str(e)}")
            return {
                'success': False,
                'error': f'Textract async processing failed: {str(e)}',
                'text': ''
            }
    
    def _analyze_with_comprehend(self, text: str) -> Dict[str, Any]:
        """Analyze text with Amazon Comprehend for entities and key phrases"""
        
        if not text or len(text.strip()) < 10:
            return {
                'entities': [],
                'key_phrases': [],
                'sentiment': None,
                'language': 'en'
            }
        
        try:
            # Truncate text if too long (Comprehend limits)
            max_text_length = 5000  # Comprehend limit
            if len(text) > max_text_length:
                text = text[:max_text_length]
                logger.info(f"Text truncated to {max_text_length} characters for Comprehend analysis")
            
            analysis_result = {
                'entities': [],
                'key_phrases': [],
                'sentiment': None,
                'language': 'en'
            }
            
            # Detect language
            try:
                language_response = self.comprehend_client.detect_dominant_language(Text=text)
                languages = language_response.get('Languages', [])
                if languages:
                    analysis_result['language'] = languages[0]['LanguageCode']
            except Exception as e:
                logger.warning(f"Language detection failed: {str(e)}")
            
            # Extract entities
            try:
                entities_response = self.comprehend_client.detect_entities(
                    Text=text,
                    LanguageCode=analysis_result['language']
                )
                
                entities = []
                for entity in entities_response.get('Entities', []):
                    entities.append({
                        'text': entity['Text'],
                        'type': entity['Type'],
                        'score': entity['Score'],
                        'begin_offset': entity['BeginOffset'],
                        'end_offset': entity['EndOffset']
                    })
                
                analysis_result['entities'] = entities
                
            except Exception as e:
                logger.warning(f"Entity detection failed: {str(e)}")
            
            # Extract key phrases
            try:
                phrases_response = self.comprehend_client.detect_key_phrases(
                    Text=text,
                    LanguageCode=analysis_result['language']
                )
                
                key_phrases = []
                for phrase in phrases_response.get('KeyPhrases', []):
                    key_phrases.append({
                        'text': phrase['Text'],
                        'score': phrase['Score'],
                        'begin_offset': phrase['BeginOffset'],
                        'end_offset': phrase['EndOffset']
                    })
                
                analysis_result['key_phrases'] = key_phrases
                
            except Exception as e:
                logger.warning(f"Key phrase detection failed: {str(e)}")
            
            # Analyze sentiment
            try:
                sentiment_response = self.comprehend_client.detect_sentiment(
                    Text=text,
                    LanguageCode=analysis_result['language']
                )
                
                analysis_result['sentiment'] = {
                    'sentiment': sentiment_response['Sentiment'],
                    'scores': sentiment_response['SentimentScore']
                }
                
            except Exception as e:
                logger.warning(f"Sentiment analysis failed: {str(e)}")
            
            logger.info(f"Comprehend analysis completed: {len(analysis_result['entities'])} entities, {len(analysis_result['key_phrases'])} key phrases")
            return analysis_result
            
        except Exception as e:
            logger.error(f"Comprehend analysis failed: {str(e)}")
            return {
                'entities': [],
                'key_phrases': [],
                'sentiment': None,
                'language': 'en',
                'error': str(e)
            }
    
    def _extract_from_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Only add non-empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_content.append(' | '.join(row_text))
            
            # Combine all content
            all_content = []
            if paragraphs:
                all_content.extend(paragraphs)
            if table_content:
                all_content.append('\n--- Tables ---')
                all_content.extend(table_content)
            
            full_text = '\n\n'.join(all_content)
            
            return {
                'success': True,
                'text': full_text,
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'table_rows_extracted': len(table_content)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Failed to extract text from DOCX: {str(e)}',
                'text': ''
            }
    
    def validate_extracted_text(self, text: str, min_length: int = 10) -> Dict[str, Any]:
        """
        Validate extracted text quality
        
        Args:
            text: Extracted text to validate
            min_length: Minimum acceptable text length
            
        Returns:
            Validation results
        """
        
        validation_result = {
            'is_valid': True,
            'warnings': [],
            'statistics': {}
        }
        
        # Basic statistics
        validation_result['statistics'] = {
            'character_count': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.splitlines()),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()])
        }
        
        # Validation checks
        if len(text) < min_length:
            validation_result['is_valid'] = False
            validation_result['warnings'].append(f'Text too short: {len(text)} characters (minimum: {min_length})')
        
        # Check for mostly non-alphabetic content (might indicate extraction issues)
        alphabetic_chars = sum(1 for c in text if c.isalpha())
        if len(text) > 0:
            alphabetic_ratio = alphabetic_chars / len(text)
            if alphabetic_ratio < 0.3:
                validation_result['warnings'].append(f'Low alphabetic content ratio: {alphabetic_ratio:.2f}')
        
        # Check for excessive whitespace or special characters
        whitespace_ratio = sum(1 for c in text if c.isspace()) / len(text) if len(text) > 0 else 0
        if whitespace_ratio > 0.7:
            validation_result['warnings'].append(f'Excessive whitespace: {whitespace_ratio:.2f}')
        
        return validation_result
    
    def clean_extracted_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.splitlines():
            cleaned_line = ' '.join(line.split())  # Normalize whitespace
            if cleaned_line:  # Only keep non-empty lines
                lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive newlines (more than 2 consecutive)
        import re
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()


# Global instance for easy import
text_extractor = TextExtractor()


# Helper functions for common operations

def extract_text_from_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Simple helper function to extract text from file
    
    Args:
        file_content: Raw file content as bytes
        filename: Original filename
        
    Returns:
        Extracted text or None if extraction failed
    """
    
    result = text_extractor.extract_text(file_content, filename)
    
    if result['success']:
        return text_extractor.clean_extracted_text(result['text'])
    else:
        logger.error(f"Text extraction failed for {filename}: {result['error']}")
        return None


def get_supported_file_types() -> Dict[str, str]:
    """Get dictionary of supported file types and descriptions"""
    
    return {
        '.pdf': 'Portable Document Format',
        '.docx': 'Microsoft Word Document',
        '.txt': 'Plain Text File'
    }


def is_text_extraction_available() -> bool:
    """Check if text extraction libraries are available"""
    return TEXT_EXTRACTION_AVAILABLE