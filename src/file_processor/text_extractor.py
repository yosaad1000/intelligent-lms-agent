"""
Text Extraction Engine for File Processor Microservice

This module handles text extraction from various file formats including
PDF, DOCX, and TXT files with advanced processing capabilities.
"""

import os
import re
import logging
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import tempfile

# PDF processing libraries
import PyPDF2
import pdfplumber

# DOCX processing library
from docx import Document

# Text processing
import unicodedata

logger = logging.getLogger(__name__)

class TextExtractionError(Exception):
    """Base exception for text extraction errors"""
    pass

class PDFExtractionError(TextExtractionError):
    """PDF-specific extraction error"""
    pass

class DOCXExtractionError(TextExtractionError):
    """DOCX-specific extraction error"""
    pass

class TextProcessingError(TextExtractionError):
    """Text processing error"""
    pass

class ExtractionResult:
    """Result object for text extraction operations"""
    
    def __init__(self, success: bool, text: str = "", metadata: Dict[str, Any] = None,
                 error: str = None, extraction_method: str = ""):
        self.success = success
        self.text = text
        self.metadata = metadata or {}
        self.error = error
        self.extraction_method = extraction_method
        self.word_count = len(text.split()) if text else 0
        self.char_count = len(text) if text else 0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert result to dictionary"""
        return {
            'success': self.success,
            'text': self.text,
            'metadata': self.metadata,
            'error': self.error,
            'extraction_method': self.extraction_method,
            'word_count': self.word_count,
            'char_count': self.char_count
        }

class AdvancedTextExtractor:
    """
    Advanced text extraction engine supporting multiple file formats
    with robust error handling and content analysis.
    """
    
    def __init__(self):
        """Initialize the text extractor"""
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        logger.info("AdvancedTextExtractor initialized")
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return ExtractionResult(
                    success=False,
                    error=f"File does not exist: {file_path}"
                )
            
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return self.extract_pdf_text(str(file_path))
            elif extension == '.docx':
                return self.extract_docx_text(str(file_path))
            elif extension == '.txt':
                return self.extract_txt_text(str(file_path))
            else:
                return ExtractionResult(
                    success=False,
                    error=f"Unsupported file extension: {extension}"
                )
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return ExtractionResult(
                success=False,
                error=f"Extraction failed: {e}"
            )
    
    def extract_pdf_text(self, file_path: str) -> ExtractionResult:
        """
        Extract text from PDF files using multiple methods for robustness.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            pdfplumber_result = self._extract_pdf_with_pdfplumber(file_path)
            if pdfplumber_result.success and pdfplumber_result.text.strip():
                return pdfplumber_result
            
            # Fallback to PyPDF2
            pypdf2_result = self._extract_pdf_with_pypdf2(file_path)
            if pypdf2_result.success and pypdf2_result.text.strip():
                return pypdf2_result
            
            # If both methods failed or returned empty text
            if pdfplumber_result.error and pypdf2_result.error:
                return ExtractionResult(
                    success=False,
                    error=f"Both extraction methods failed. PDFPlumber: {pdfplumber_result.error}, PyPDF2: {pypdf2_result.error}"
                )
            
            # Return the result with more text, even if minimal
            if len(pdfplumber_result.text) >= len(pypdf2_result.text):
                return pdfplumber_result
            else:
                return pypdf2_result
                
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ExtractionResult(
                success=False,
                error=f"PDF extraction error: {e}"
            )
    
    def _extract_pdf_with_pdfplumber(self, file_path: str) -> ExtractionResult:
        """Extract PDF text using pdfplumber (better for complex layouts)"""
        try:
            text_content = []
            tables_content = []
            metadata = {
                'pages': 0,
                'tables_found': 0,
                'extraction_method': 'pdfplumber'
            }
            
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['tables_found'] += len(tables)
                        for table_num, table in enumerate(tables, 1):
                            table_text = self._format_table_as_text(table)
                            tables_content.append(f"--- Page {page_num}, Table {table_num} ---\n{table_text}")
            
            # Combine text and tables
            full_text = "\n\n".join(text_content)
            if tables_content:
                full_text += "\n\n=== TABLES ===\n\n" + "\n\n".join(tables_content)
            
            return ExtractionResult(
                success=True,
                text=full_text,
                metadata=metadata,
                extraction_method="pdfplumber"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"PDFPlumber extraction failed: {e}",
                extraction_method="pdfplumber"
            )
    
    def _extract_pdf_with_pypdf2(self, file_path: str) -> ExtractionResult:
        """Extract PDF text using PyPDF2 (fallback method)"""
        try:
            text_content = []
            metadata = {
                'pages': 0,
                'extraction_method': 'pypdf2'
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    # Try to decrypt with empty password
                    try:
                        pdf_reader.decrypt("")
                    except:
                        return ExtractionResult(
                            success=False,
                            error="PDF is encrypted and cannot be decrypted",
                            extraction_method="pypdf2"
                        )
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
            
            full_text = "\n\n".join(text_content)
            
            return ExtractionResult(
                success=True,
                text=full_text,
                metadata=metadata,
                extraction_method="pypdf2"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"PyPDF2 extraction failed: {e}",
                extraction_method="pypdf2"
            )
    
    def extract_docx_text(self, file_path: str) -> ExtractionResult:
        """
        Extract text from DOCX files with formatting preservation.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            text_content = []
            metadata = {
                'paragraphs': 0,
                'tables': 0,
                'headers': 0,
                'footers': 0,
                'extraction_method': 'python-docx'
            }
            
            doc = Document(file_path)
            
            # Extract headers
            for section in doc.sections:
                header = section.header
                if header:
                    header_text = self._extract_docx_element_text(header)
                    if header_text.strip():
                        text_content.append(f"=== HEADER ===\n{header_text}")
                        metadata['headers'] += 1
            
            # Extract main document paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_content.append(para_text)
                    metadata['paragraphs'] += 1
            
            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = self._extract_docx_table_text(table)
                if table_text.strip():
                    text_content.append(f"=== TABLE {table_num} ===\n{table_text}")
                    metadata['tables'] += 1
            
            # Extract footers
            for section in doc.sections:
                footer = section.footer
                if footer:
                    footer_text = self._extract_docx_element_text(footer)
                    if footer_text.strip():
                        text_content.append(f"=== FOOTER ===\n{footer_text}")
                        metadata['footers'] += 1
            
            full_text = "\n\n".join(text_content)
            
            return ExtractionResult(
                success=True,
                text=full_text,
                metadata=metadata,
                extraction_method="python-docx"
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ExtractionResult(
                success=False,
                error=f"DOCX extraction failed: {e}",
                extraction_method="python-docx"
            )
    
    def extract_txt_text(self, file_path: str) -> ExtractionResult:
        """
        Extract text from TXT files with encoding detection.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        
                    metadata = {
                        'encoding': encoding,
                        'lines': len(text.splitlines()),
                        'extraction_method': 'text-file'
                    }
                    
                    return ExtractionResult(
                        success=True,
                        text=text,
                        metadata=metadata,
                        extraction_method="text-file"
                    )
                    
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Failed to read with encoding {encoding}: {e}")
                    continue
            
            # If all encodings failed
            return ExtractionResult(
                success=False,
                error="Could not decode text file with any supported encoding",
                extraction_method="text-file"
            )
            
        except Exception as e:
            logger.error(f"TXT extraction failed for {file_path}: {e}")
            return ExtractionResult(
                success=False,
                error=f"TXT extraction failed: {e}",
                extraction_method="text-file"
            )
    
    def _format_table_as_text(self, table: List[List[str]]) -> str:
        """Format a table as readable text"""
        if not table:
            return ""
        
        # Find maximum width for each column
        col_widths = []
        for row in table:
            for i, cell in enumerate(row or []):
                cell_text = str(cell) if cell else ""
                if i >= len(col_widths):
                    col_widths.append(len(cell_text))
                else:
                    col_widths[i] = max(col_widths[i], len(cell_text))
        
        # Format table
        formatted_rows = []
        for row in table:
            if row:
                formatted_cells = []
                for i, cell in enumerate(row):
                    cell_text = str(cell) if cell else ""
                    width = col_widths[i] if i < len(col_widths) else 0
                    formatted_cells.append(cell_text.ljust(width))
                formatted_rows.append(" | ".join(formatted_cells))
        
        return "\n".join(formatted_rows)
    
    def _extract_docx_element_text(self, element) -> str:
        """Extract text from a DOCX element (header/footer)"""
        text_parts = []
        for paragraph in element.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        return "\n".join(text_parts)
    
    def _extract_docx_table_text(self, table) -> str:
        """Extract text from a DOCX table"""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            if any(row_text):  # Only add non-empty rows
                table_text.append(" | ".join(row_text))
        return "\n".join(table_text)
    
    def clean_and_normalize_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        try:
            # Normalize Unicode characters
            text = unicodedata.normalize('NFKD', text)
            
            # Remove excessive whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # Remove excessive line breaks
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            
            # Clean up common PDF artifacts
            text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\\\n\r]', '', text)
            
            # Remove page numbers and headers/footers patterns
            text = re.sub(r'\n\s*\d+\s*\n', '\n', text)  # Standalone page numbers
            text = re.sub(r'\n\s*Page \d+.*?\n', '\n', text, flags=re.IGNORECASE)
            
            # Strip and return
            return text.strip()
            
        except Exception as e:
            logger.warning(f"Text cleaning failed: {e}")
            return text.strip()
    
    def chunk_text_for_indexing(self, text: str, max_chunk_size: int = 1000, 
                               overlap: int = 100) -> List[Dict[str, Any]]:
        """
        Chunk text into optimal sizes for indexing with overlap.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks with metadata
        """
        if not text:
            return []
        
        try:
            # Split into sentences for better chunking
            sentences = re.split(r'(?<=[.!?])\s+', text)
            
            chunks = []
            current_chunk = ""
            current_size = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                
                # If adding this sentence would exceed the limit
                if current_size + sentence_size > max_chunk_size and current_chunk:
                    # Save current chunk
                    chunks.append({
                        'text': current_chunk.strip(),
                        'size': current_size,
                        'chunk_index': len(chunks),
                        'word_count': len(current_chunk.split())
                    })
                    
                    # Start new chunk with overlap
                    if overlap > 0 and current_chunk:
                        overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                        current_chunk = overlap_text + " " + sentence
                        current_size = len(current_chunk)
                    else:
                        current_chunk = sentence
                        current_size = sentence_size
                else:
                    # Add sentence to current chunk
                    if current_chunk:
                        current_chunk += " " + sentence
                        current_size += sentence_size + 1  # +1 for space
                    else:
                        current_chunk = sentence
                        current_size = sentence_size
            
            # Add final chunk if it exists
            if current_chunk.strip():
                chunks.append({
                    'text': current_chunk.strip(),
                    'size': len(current_chunk),
                    'chunk_index': len(chunks),
                    'word_count': len(current_chunk.split())
                })
            
            return chunks
            
        except Exception as e:
            logger.error(f"Text chunking failed: {e}")
            # Return single chunk as fallback
            return [{
                'text': text,
                'size': len(text),
                'chunk_index': 0,
                'word_count': len(text.split())
            }]
    
    def extract_concepts_and_keywords(self, text: str, max_concepts: int = 20) -> List[str]:
        """
        Extract key concepts and keywords from text.
        
        Args:
            text: Text to analyze
            max_concepts: Maximum number of concepts to return
            
        Returns:
            List of extracted concepts
        """
        if not text:
            return []
        
        try:
            # Simple keyword extraction based on frequency and length
            # Remove common stop words
            stop_words = {
                'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have',
                'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
                'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we',
                'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'her',
                'its', 'our', 'their', 'page', 'table', 'header', 'footer'
            }
            
            # Extract words (3+ characters, alphanumeric)
            words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
            
            # Filter out stop words and count frequency
            word_freq = {}
            for word in words:
                if word not in stop_words and len(word) >= 3:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Sort by frequency and length (prefer longer, more frequent words)
            concepts = sorted(
                word_freq.items(),
                key=lambda x: (x[1], len(x[0])),  # Sort by frequency, then length
                reverse=True
            )
            
            # Return top concepts
            return [concept[0] for concept in concepts[:max_concepts]]
            
        except Exception as e:
            logger.warning(f"Concept extraction failed: {e}")
            return []

def main():
    """Test function for AdvancedTextExtractor"""
    logging.basicConfig(level=logging.INFO)
    
    print("🧪 Testing AdvancedTextExtractor")
    print("=" * 50)
    
    extractor = AdvancedTextExtractor()
    
    # Test with a simple text file
    test_text = """
    This is a test document for text extraction.
    It contains multiple sentences and paragraphs.
    
    The purpose is to verify that our text extraction
    and processing capabilities work correctly.
    """
    
    # Test text cleaning
    cleaned = extractor.clean_and_normalize_text(test_text)
    print(f"✅ Text cleaning: {len(cleaned)} characters")
    
    # Test chunking
    chunks = extractor.chunk_text_for_indexing(cleaned, max_chunk_size=100)
    print(f"✅ Text chunking: {len(chunks)} chunks created")
    
    # Test concept extraction
    concepts = extractor.extract_concepts_and_keywords(cleaned)
    print(f"✅ Concept extraction: {len(concepts)} concepts found")
    print(f"   Top concepts: {', '.join(concepts[:5])}")
    
    print("\n🎉 AdvancedTextExtractor test completed!")

if __name__ == "__main__":
    main()